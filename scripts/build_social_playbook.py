from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "strategy"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "Haven_Khareef_Social_Media_Playbook_2026.docx"
LOGO = ROOT / "assets" / "img" / "haven-logo-original.jpg"

NAVY = "082338"
NAVY_2 = "12354A"
INK = "17232B"
MUTED = "66747C"
SAND = "E9DFC9"
SAND_2 = "F5F0E7"
FOG = "DDE7E5"
CORAL = "A33B32"
WHITE = "FFFFFF"
LINE = "C9D2D5"

FONT = "Aptos"
ARABIC_FONT = "Arial"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color=INK, italic=False, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), ARABIC_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def add_hyperlink(paragraph, text: str, url: str, color=NAVY_2, size=8.6):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(round(size * 2))))
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(round(size * 2))))
    r_pr.extend([r_fonts, color_el, underline, sz, sz_cs])
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.extend([r_pr, text_el])
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def make_num(fmt: str, text: str, left: int, hanging: int):
        nonlocal next_abs, next_num
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abs))
        mlt = OxmlElement("w:multiLevelType")
        mlt.set(qn("w:val"), "singleLevel")
        abstract.append(mlt)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(next_abs))
        num.append(abs_id)
        numbering.append(num)
        result = next_num
        next_abs += 1
        next_num += 1
        return result

    return {
        "bullet": make_num("bullet", "•", 540, 270),
        "number": make_num("decimal", "%1.", 540, 270),
    }


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def clone_num_id(doc: Document, base_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    base = None
    for node in numbering.findall(qn("w:num")):
        if int(node.get(qn("w:numId"))) == base_num_id:
            base = node
            break
    if base is None:
        return base_num_id
    abstract = base.find(qn("w:abstractNumId"))
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract_copy = OxmlElement("w:abstractNumId")
    abstract_copy.set(qn("w:val"), abstract.get(qn("w:val")))
    num.append(abstract_copy)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return new_id


def style_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.86)
    sec.right_margin = Inches(0.86)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), ARABIC_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, before, after, color in [
        ("Title", 28, 0, 6, NAVY),
        ("Subtitle", 13, 0, 12, MUTED),
        ("Heading 1", 18, 16, 7, NAVY),
        ("Heading 2", 13.5, 12, 5, NAVY_2),
        ("Heading 3", 11.5, 9, 4, CORAL),
    ]:
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st._element.rPr.rFonts.set(qn("w:cs"), ARABIC_FONT)
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("HAVEN  /  KHAREEF 2026 SOCIAL PLAYBOOK"), 8, True, MUTED)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("OSARA, SALALAH   •   "), 8, False, MUTED)
    add_field(p, "PAGE")


def add_para(doc, text="", bold_lead=None, size=10.5, color=INK, align=None, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), size, True, color, italic)
        set_run(p.add_run(text[len(bold_lead):]), size, False, color, italic)
    else:
        set_run(p.add_run(text), size, False, color, italic)
    return p


def add_bullets(doc, items, num_id, compact=False):
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(3 if compact else 5)
        if isinstance(item, tuple):
            label, text = item
            set_run(p.add_run(label), 10.25, True, NAVY)
            set_run(p.add_run(text), 10.25, False, INK)
        else:
            set_run(p.add_run(item), 10.25)


def add_steps(doc, items, num_id):
    list_num_id = clone_num_id(doc, num_id)
    for title, text in items:
        p = doc.add_paragraph()
        apply_num(p, list_num_id)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(6)
        set_run(p.add_run(title + " — "), 10.25, True, NAVY)
        set_run(p.add_run(text), 10.25)


def add_callout(doc, label, text, fill=SAND_2, accent=CORAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    set_run(p.add_run(label.upper() + "  "), 9, True, accent)
    set_run(p.add_run(text), 10.3, True, NAVY)
    return p


def add_table(doc, headers, rows, widths, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths, indent_dxa=120)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(head), 9, True, WHITE)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if r_idx % 2:
                set_cell_shading(cells[i], "F7F9F9")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), 8.7, False, INK)
    set_table_geometry(table, widths, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def section_break(doc):
    doc.add_page_break()


def add_source(doc, title, url, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_hyperlink(p, title, url)
    set_run(p.add_run(" — " + note), 8.6, False, MUTED)


doc = Document()
style_document(doc)
numbering = configure_numbering(doc)
for section in doc.sections:
    add_header_footer(section)

# Cover
sec = doc.sections[0]
sec.different_first_page_header_footer = True
doc.add_paragraph().paragraph_format.space_after = Pt(18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(LOGO), width=Inches(4.65))
doc_pr = p._p.xpath(".//wp:docPr")[0]
doc_pr.set("title", "Haven logo")
doc_pr.set("descr", "Haven bilingual logo with a traditional sailing dhow on a navy background.")
doc.add_paragraph().paragraph_format.space_after = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("THE HAVEN EFFECT"), 28, True, NAVY)
p.paragraph_format.space_after = Pt(4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Khareef Social Media Instruction Book"), 16, False, NAVY_2)
p.paragraph_format.space_after = Pt(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Instagram • Snapchat • Local Discovery • Footfall"), 11, True, CORAL)
p.paragraph_format.space_after = Pt(22)
add_callout(
    doc,
    "Operating objective",
    "Make Haven the beachfront plan people send to one another today—and make the next step to visit unmistakably easy.",
    fill=SAND,
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Prepared for Haven, Osara, Salalah  |  25 July 2026"), 9.5, False, MUTED)

section_break(doc)

# Executive directive
doc.add_heading("Executive directive", level=1)
add_callout(
    doc,
    "The strategy in one sentence",
    "Do not market a café with a beach nearby. Market a scarce Khareef moment—sea, mist, seating, company and a Haven signature in hand—then remove every obstacle between discovery and arrival.",
)
add_para(
    doc,
    "Haven’s strongest awareness asset is not a generic cup shot. It is the combination competitors cannot copy exactly: a beachfront seat in this week’s Khareef atmosphere, served with a credible premium café product. The creative system should therefore open with place or sensation, prove the product, show real people enjoying it, and end with a location-and-time instruction.",
)
doc.add_heading("What must happen this week", level=2)
add_steps(
    doc,
    [
        ("Build the conversion surface", "Make the bio, map link, hours, price/menu and message route correct before pushing reach."),
        ("Capture one concentrated media bank", "Shoot enough vertical footage in one dawn and one golden-hour session to publish for seven days."),
        ("Launch three repeatable series", "“Your Khareef seat”, “Made at Haven”, and “Tonight at the beach” create memory through repetition."),
        ("Recruit distribution, not fame", "Use local micro-creators, café guides and guests whose audiences are already in Salalah or travelling now."),
        ("Measure arrivals", "Use a POS source button or verbal code, direction clicks, DMs and redemptions—not views alone."),
    ],
    numbering["number"],
)
doc.add_heading("The commercial guardrail", level=2)
add_para(
    doc,
    "Virality is useful only when it reaches plausible visitors and preserves the experience shown. A Reel with 30,000 local/GCC views, clear directions and 80 measured visits is more valuable than 500,000 distant views with no arrival path. Protect service speed, seating truth and product consistency; content amplifies both strengths and failures.",
)

doc.add_heading("Recommended content allocation", level=2)
add_table(
    doc,
    ["Job", "Share", "What it shows", "Primary signal"],
    [
        ("Place desire", "35%", "Beach, mist, seating, light, people", "Sends / shares"),
        ("Product desire", "25%", "Signature cold brew, Spanish latte, brownie", "Watch time / saves"),
        ("Social proof", "20%", "Guests, creators, reactions, tagged content", "Profile visits / replies"),
        ("Useful planning", "15%", "Hours, location, parking, availability", "Directions / DMs"),
        ("Direct offer", "5%", "Time-bound bundle or reason to act now", "Redemption / visits"),
    ],
    [1700, 900, 4160, 2600],
)

section_break(doc)

# Market
doc.add_heading("1. Market reality: why Instagram and Snapchat lead", level=1)
add_para(
    doc,
    "Oman is a high-connectivity, mobile-first market. DataReportal reported 5.14 million internet users and 3.29 million social-media user identities at the start of 2025. Platform figures below are advertising-reach estimates—not monthly active users—and should be used directionally.",
)
add_table(
    doc,
    ["Platform", "Oman ad reach", "Adult/eligible context", "Implication for Haven"],
    [
        ("Instagram", "2.50m", "61.8% of adults 18+", "Primary discovery, visual brand and creator collaboration."),
        ("Snapchat", "2.15m", "48.6% of adults 18+; +12.4% YoY reach", "Primary immediacy, young/GCC audience and live proof."),
        ("TikTok", "1.83m adults", "47.1% of adults 18+", "Secondary discovery engine; repurpose only native edits."),
        ("YouTube", "3.29m", "64.0% of internet users", "Shorts and search shelf-life; low-cost reuse."),
    ],
    [1450, 1550, 2700, 3660],
)
add_para(
    doc,
    "Khareef concentrates demand. Official preliminary statistics counted 827,115 visitors from 21 June to 15 August 2025; 71.5% were Omani, 17.3% other GCC, and 11.1% other nationalities. Nearly 46.5% arrived in the first half of August. The full 2024 season reached about 1.048 million visitors. This is a seasonal attention market: timely local distribution matters more than building a distant audience slowly.",
)
add_callout(
    doc,
    "What this means",
    "Create in Arabic first or genuinely bilingual; optimise for people already in Dhofar and GCC travellers planning within days; publish live atmosphere and availability because the product expires with the moment.",
    fill=FOG,
)

doc.add_heading("Audience jobs-to-be-done", level=2)
add_bullets(
    doc,
    [
        ("Families: ", "“Give us a comfortable, attractive place where the outing feels easy and worth the drive.”"),
        ("Friends/couples: ", "“Give us a plan worth sending to the group and a scene that feels special tonight.”"),
        ("Coffee-led guests: ", "“Prove the drink is more than scenery; show craft, portion and flavour.”"),
        ("Tourists: ", "“Help us find it, know when to go, and avoid uncertainty about hours, seating and price.”"),
        ("Local repeat guests: ", "“Give us a new reason to return this week—a ritual, limited pairing, moment or guest feature.”"),
    ],
    numbering["bullet"],
)

section_break(doc)

# Benchmark
doc.add_heading("2. Comparable-business study", level=1)
add_para(
    doc,
    "Method: public profiles and visible recent posts were inspected on 25 July 2026, supplemented by official sites, press coverage and customer-review themes. Follower counts are snapshots, not proof that a specific tactic caused growth. The useful evidence is the repeated operating pattern.",
)
add_table(
    doc,
    ["Business", "Visible pattern", "Why it works", "Haven adaptation"],
    [
        ("Voliere, Salalah\n17.1K IG", "Exact hours/locations; recent opening content; own Reels mixed with café-guide and guest collaborations.", "Converts novelty into distributed local proof.", "Collab-post the opening/Khareef moment with local guides; pin exact operating facts."),
        ("Lantana, Al Baleed\n8.2K IG", "Beach view; breakfast/daypart; creator guest posts; Highlights for Menu, Location and Timing.", "Sells an occasion and reduces planning friction.", "Create Start Here, Location, Menu, Seating and Guests Highlights immediately."),
        ("55 Coffee, Oman\n62.2K IG", "One recognisable master brand; branch link; seasonal product creative; mix of Reels and carousels.", "Scale comes from consistency and product memory.", "Repeat one signature drink and one visual code until remembered."),
        ("SALT, UAE", "Iconic beach asset; simple hero product; community/movement language; pop-ups and surprise.", "People share identity and belonging, not only food.", "Own “your Khareef seat” as a social ritual, not a furniture rental."),
        ("% Arabica, UAE", "Minimal product/worldview; coffee rituals; guest-led ideas; simple questions and direct order CTA.", "A disciplined point of view builds recognition.", "Keep premium feed restrained; put urgency and messier human proof in Stories/Snap."),
        ("DRIFT / Koko Bay", "Destination imagery, day-to-night programming, offers, availability and booking routes.", "Monetises scenery by packaging occasions.", "Publish time-specific plans: morning calm, sunset mist, after-dinner coffee."),
    ],
    [1500, 2800, 2200, 2860],
)

doc.add_heading("Distilled winning process", level=2)
add_steps(
    doc,
    [
        ("Own one visual territory", "A recognisable location, colour, object or ritual gives every post memory value."),
        ("Turn atmosphere into a plan", "Name the occasion and daypart instead of describing the brand abstractly."),
        ("Borrow trusted distribution", "Use Collab posts, tagged guests and local niche pages so discovery arrives with social proof."),
        ("Make utility visible", "Hours, map, menu, price cues and seating truth belong in highlights and the final frame."),
        ("Repeat the hero", "One signature item and one recurring series outperform a constantly changing set of unrelated ideas."),
        ("Close the loop", "Track the source at the counter and feed real guest reactions back into content within 24 hours."),
    ],
    numbering["number"],
)
add_callout(
    doc,
    "Competitive opening",
    "Haven can be more human and more useful than polished competitors while still looking premium. The white space is premium editorial craft plus live beachfront truth.",
)

section_break(doc)

# Funnel
doc.add_heading("3. The awareness-to-arrival system", level=1)
add_table(
    doc,
    ["Stage", "Audience question", "Content answer", "Metric"],
    [
        ("STOP", "Why should I pause?", "Mist on the beach, first-pour motion, human expression, direct local hook.", "3-sec hold / swipe-away"),
        ("WANT", "Why Haven?", "Seat + sea + signature product + emotional occasion.", "Completion / rewatch / saves"),
        ("SEND", "Who would come with me?", "Group-plan language, relatable moment, local novelty.", "Shares / sends"),
        ("PLAN", "Where, when, how much?", "Location, hours, availability, menu and parking cues.", "Profile / map / DM"),
        ("ARRIVE", "Why now?", "Live atmosphere, limited moment, simple CTA or approved offer.", "POS source / redemption"),
        ("AMPLIFY", "Was it worth it?", "Fast service, product truth, guest prompt and review request.", "UGC / reviews / repeat"),
    ],
    [1150, 2200, 4050, 1960],
)
doc.add_heading("The two equations", level=2)
add_callout(
    doc,
    "Shareability",
    "First-frame clarity × sensory/emotional payoff × send-to-a-friend relevance × local novelty × completion.",
    fill=FOG,
)
add_callout(
    doc,
    "Conversion",
    "Qualified reach × desire × distance/availability clarity × social proof × operational trust × measurable CTA.",
    fill=SAND_2,
)
add_para(
    doc,
    "A post does not need every component equally. Reels maximise STOP, WANT and SEND; Stories and Snapchat maximise PLAN and ARRIVE; Google Maps converts high intent; service and guest prompts create AMPLIFY.",
)

doc.add_heading("Five content pillars", level=2)
add_bullets(
    doc,
    [
        ("1. Your Khareef seat (35%): ", "the sea, mist, seating, weather, arrival and changing light."),
        ("2. Made at Haven (25%): ", "real preparation, barista hands, product truth, texture and founder recipe."),
        ("3. People make the place (20%): ", "guests, staff warmth, group moments, testimonials and tagged content."),
        ("4. Useful now (15%): ", "today’s hours, live seating, map, weather fit, parking and what to order."),
        ("5. The reason this week (5%): ", "a limited pairing, small event, guest barista, sunset service or approved bundle."),
    ],
    numbering["bullet"],
)

section_break(doc)

# Instagram
doc.add_heading("4. Instagram instruction book", level=1)
doc.add_heading("Profile conversion checklist", level=2)
add_bullets(
    doc,
    [
        "Name field: “Haven Beach Café • Salalah” (searchable, not only the brand name).",
        "Bio line 1: premium café + beachfront seating + Osara, Salalah.",
        "Bio line 2: today’s opening hours; update immediately if weather or operations change.",
        "One link page with Google Maps first, menu second, WhatsApp/message third.",
        "Action buttons: directions/contact if available. Never bury the map under several promotional links.",
        "Pinned posts: (1) the Haven experience, (2) how to find/when to visit, (3) signature menu proof.",
        "Highlights: START HERE / LOCATION / MENU / SEATING / GUESTS / TODAY. Use bilingual covers.",
        "Check Account Status and recommendation eligibility before the campaign.",
    ],
    numbering["bullet"],
    compact=True,
)

doc.add_heading("Format roles and cadence for this week", level=2)
add_table(
    doc,
    ["Format", "Cadence", "Ideal length", "Job"],
    [
        ("Reels", "1 strong post daily; 2 only if distinct", "7–15 sec discovery; 20–35 sec story", "Cold reach, shares, creator distribution."),
        ("Stories", "8–15 frames across the day", "1 idea per 5–7 sec frame", "Live proof, FAQs, map, availability, replies."),
        ("Carousel", "2 this week", "5–7 slides", "Saveable planning: menu, “best time”, what to expect."),
        ("Collab post", "3–5 this week", "Use creator’s strongest native cut", "Combine audiences without reposting."),
        ("Live", "Only with a real moment", "5–12 min", "Sunset, event or Q&A—not empty broadcasting."),
    ],
    [1250, 1850, 2200, 4060],
)

doc.add_heading("The Haven Reel recipe", level=2)
add_steps(
    doc,
    [
        ("0.0–1.0 sec: local hook", "Show the beach/mist/hero product immediately. Text: “Your Khareef seat is here.”"),
        ("1.0–4.0 sec: sensory proof", "Condensation, pour, brownie break, feet in sand, chair facing the sea."),
        ("4.0–8.0 sec: human payoff", "A guest settles in, smiles, hands a drink to a friend or reacts genuinely."),
        ("8.0–12.0 sec: product + brand", "Full truthful serving; Haven mark visible naturally, not as a long logo card."),
        ("Final 1.5 sec: action", "“Osara, Salalah • open today • map in bio • send this to your person.”"),
    ],
    numbering["number"],
)
add_para(
    doc,
    "Technical baseline: 9:16, 1080×1920 master, 24/30 fps, no important text near the top/bottom UI zones, captions burned in, clean natural sound plus licensed/native audio. Shoot every key action wide, medium and macro. Do not export with another platform’s watermark.",
)

doc.add_heading("Hooks that fit Haven", level=2)
add_bullets(
    doc,
    [
        "“The Khareef plan you send before someone else finds the seats.”",
        "“POV: your coffee comes with this view.”",
        "“Salalah is 22°C. Your seat is facing the sea.” (use only with verified live weather).",
        "“What 2.4 OMR looks like at the beach: Haven Passion Fruit Cold Brew.”",
        "“Don’t order the usual. The founding barista made this for Khareef.”",
        "“From empty beach to tonight’s Haven in 8 seconds.”",
        "“Save this for your next Osara drive.”",
        "“The brownie break that tells you it is not just the view.”",
    ],
    numbering["bullet"],
    compact=True,
)

doc.add_heading("Stories: the daily conversion ladder", level=2)
add_steps(
    doc,
    [
        ("Morning", "Weather/sea proof + opening time + calm seating. Add Location sticker."),
        ("Pre-peak", "Hero product + menu price + poll: cold brew or Spanish latte?"),
        ("Arrival window", "Short route landmark/parking cue + map link + honest seat status."),
        ("Peak", "Real guest atmosphere + staff action. Avoid implying empty seats if full."),
        ("Late", "Last-order time + warm drink/dessert + repost tagged guests."),
        ("Close", "Thank guests, answer top FAQ, save useful frames to Highlights."),
    ],
    numbering["number"],
)

doc.add_heading("Caption and community rules", level=2)
add_bullets(
    doc,
    [
        "Arabic first when the post is local-action oriented; concise English beneath it for GCC/international visitors.",
        "First line names the plan or payoff. Do not begin with a generic slogan.",
        "Use 3–6 precise tags/keywords, not a cloud: Haven, Salalah, Khareef, Osara, beachfront café, signature drink.",
        "Geotag every location-led post. Add collaborator instead of downloading/reuploading creator work.",
        "Reply to visit-intent comments and DMs within 15 minutes during campaign windows.",
        "Turn repeated questions into the next Story. Pin useful comments.",
        "Ask for the send: “Send this to who is coming with you,” not “please share.”",
    ],
    numbering["bullet"],
)

doc.add_heading("Paid Instagram recommendation", level=2)
add_para(
    doc,
    "Do not boost everything. After 2–4 hours, promote the organic Reel with the best qualified share rate and hold. Run two ad sets: (A) people currently in/near Salalah; (B) GCC travellers in Dhofar or an approved travel-interest audience. Use Awareness/video-view optimisation for discovery and a separate Traffic/message campaign for map/WhatsApp actions. Keep the best 9:16 creative, sound and safe-zone text; Meta reports lower cost per result for native 9:16 Reels creative than image ads in its tests.",
)

section_break(doc)

# Snapchat
doc.add_heading("5. Snapchat instruction book", level=1)
add_callout(
    doc,
    "Platform role",
    "Instagram is the polished discovery shelf; Snapchat is the live evidence that Haven is worth going to now.",
)
doc.add_heading("Organic Story operating system", level=2)
add_bullets(
    doc,
    [
        "Post in short bursts across the day, not one 30-Snap dump.",
        "Open each burst with the strongest fact: sea/mist, today’s hours, available seats or the hero drink.",
        "Use phone-native footage, real voices and ambient sound. A controlled DIY feel is an advantage.",
        "Keep most Snaps 3–7 seconds. One point of interest per frame.",
        "Put the offer or urgency in the first two seconds; Snapchat recommends roughly 5–6 seconds for native-feeling ads.",
        "Design sound-on but make the message understandable visually. Snapchat says 64% of ads are viewed with sound on.",
        "Save the best public-safe series to the Public Profile; use Spotlight for the most entertaining/complete vertical cut.",
    ],
    numbering["bullet"],
)

doc.add_heading("Daily Snap sequence", level=2)
add_table(
    doc,
    ["Burst", "Frames", "Example"],
    [
        ("OPEN", "2–3", "Live beach + “Open now in Osara” + coffee machine wake-up."),
        ("PROOF", "3–5", "Signature build, brownie break, staff face, guest reaction."),
        ("UTILITY", "2–3", "Map/landmark, hours, price and seating status."),
        ("SOCIAL", "3–5", "Groups arriving, creator visit, short testimonial, repost."),
        ("LAST CALL", "2", "Late-night mood + accurate last-order time."),
    ],
    [1200, 1000, 7160],
)

doc.add_heading("Snap Ads: small local test", level=2)
add_steps(
    doc,
    [
        ("Creative A—Place", "5–6 sec: sea/mist → seat → Haven cup → “Open today, Osara.”"),
        ("Creative B—Product", "5–6 sec: passion-fruit cold brew build → price → map/visit CTA."),
        ("Creative C—UGC", "6–9 sec: creator says why they came and shows the view/product."),
        ("Target", "Begin with Salalah/Dhofar location targeting and suitable age settings; avoid over-narrow interest stacks."),
        ("Optimise", "Use Single Image or Video Ad first; retarget Public Profile/ad engagers if scale allows."),
        ("Judge", "Compare swipe-up/map/message cost and measured visits—not only completion."),
    ],
    numbering["number"],
)
add_para(
    doc,
    "Snapchat’s current specs support 9:16 creative at 720×1280 minimum; Haven should master at 1080×1920. Single Image or Video Ads can run 3–180 seconds, but short native cuts are the appropriate starting point. AR is a phase-two option only if the filter creates guest sharing—for example, subtle Khareef mist + Haven sail + location—not decorative clutter.",
)

doc.add_heading("Creator brief for Instagram + Snapchat", level=2)
add_bullets(
    doc,
    [
        "Deliverables: 1 Collab Reel (10–20 sec), 5–8 live Snaps/Stories, 3 raw vertical clips and 1 honest spoken reaction.",
        "Must show: recognisable place in first two seconds, complete product, one person enjoying it, location/hours CTA.",
        "Creator retains their voice; Haven approves factual accuracy, privacy and product truth—not every word.",
        "Request audience geography and recent Story views before agreeing. Prefer Salalah café/travel/family micro-creators over broad lifestyle reach.",
        "Track each creator with a distinct POS code or question. Disclose paid/gifted relationships using platform tools.",
    ],
    numbering["bullet"],
)

# Other channels
doc.add_heading("6. Channels that should not be missed", level=1)
doc.add_heading("Google Business Profile: highest-intent conversion surface", level=2)
add_para(
    doc,
    "Social creates the desire; Google Maps often closes the trip. Claim and fully complete the profile, use the exact pin, hours, phone, category, menu link and current exterior/route photos. Upload this week’s products and seating; ask satisfied guests for honest reviews through a QR code. Google explicitly positions Business Profile as the way to appear in Search and Maps, show offerings, collect reviews and connect with customers.",
)
add_bullets(
    doc,
    [
        "Place the Google Maps link first in the social link page.",
        "Add 10 current images: exterior approach, landmark, seating, full menu, four hero products, night appearance.",
        "Post a Khareef update and accurate special hours.",
        "Reply to every review. Never offer an incentive conditional on a positive review.",
        "Track direction requests, calls and profile actions weekly.",
    ],
    numbering["bullet"],
)

doc.add_heading("TikTok: discovery expansion", level=2)
add_para(
    doc,
    "TikTok’s adult ad reach in Oman is large enough to matter. Publish the strongest Haven concepts with TikTok-native editing, on-screen search language (“Salalah beachfront café”, “Khareef places”) and local creator voices. Do not simply upload a watermarked Instagram edit. Use 3–4 posts this week if the team can sustain replies; otherwise prioritise Instagram/Snap execution.",
)

doc.add_heading("WhatsApp: conversion layer, not a content calendar", level=2)
add_para(
    doc,
    "Use one-click WhatsApp for directions, group enquiries and live operational questions. Pre-fill a message such as “Hi Haven, are beachfront seats available around 7pm today?” Set quick replies for hours, location, menu and seating. Do not promise reservations unless operations can honour them.",
)

doc.add_heading("YouTube Shorts", level=2)
add_para(
    doc,
    "Upload 2–3 evergreen vertical cuts with searchable titles and a Maps link in the description. This is a low-effort secondary shelf because YouTube’s reported Oman reach is broad; it is not the priority community channel for this seven-day sprint.",
)

section_break(doc)

# Weekly plan
doc.add_heading("7. Seven-day Khareef launch sprint", level=1)
add_para(
    doc,
    "Dates assume immediate execution from Saturday 25 July to Friday 31 July 2026. If production starts later, preserve the sequence and shift the dates.",
)
add_table(
    doc,
    ["Day", "Hero post", "Stories / Snap", "Conversion action"],
    [
        ("Sat 25\nFoundation", "Reel: “Your Khareef seat is here” (place-led).", "Open/route/seat status; founder hello; poll.", "Fix bio, map, hours, Highlights and POS source button."),
        ("Sun 26\nSignature", "Reel: passion-fruit cold brew origin/build.", "Taste words, price, barista Q&A, reaction.", "Track “signature” orders and profile actions."),
        ("Mon 27\nUtility", "Reel: clipped 24-sec arrival route—exit, parking, gate and two areas.", "Hyperlapse cut, route landmark, area poll and FAQ.", "Pin the route; track Maps taps and ARRIVE source."),
        ("Tue 28\nPeople", "Collab Reel with local café/travel creator.", "Creator takeover burst; honest availability.", "Creator-specific code/source."),
        ("Wed 29\nSensation", "7-sec macro: condensation/pour/brownie.", "ASMR sequence; this-or-that poll.", "Promote best share-rate Reel locally."),
        ("Thu 30\nGroup plan", "Reel: “send this to tonight’s group”.", "Sunset build, arrivals, seating truth.", "Test approved pairing/limited reason."),
        ("Fri 31\nProof", "Montage of real guest moments + thanks.", "Best UGC, reviews, last-call, next-week tease.", "Retarget engagers; publish results snapshot."),
    ],
    [1050, 2750, 3350, 2210],
)

doc.add_heading("Recommended posting windows to test", level=2)
add_para(
    doc,
    "Do not treat generic “best times” as fact. Use these as operational hypotheses, then compare Haven Insights: 08:00–10:00 for morning planning/breakfast; 14:00–16:00 for evening-plan sends; 17:00–19:00 for live atmosphere; 21:00–23:00 for late café/dessert intent. Stories and Snapchat should span the actual service day.",
)

doc.add_heading("First 48-hour checklist", level=2)
add_bullets(
    doc,
    [
        "Confirm exact map pin, hours, last order, phone/WhatsApp and whether reservations are accepted.",
        "Create POS source options: IG Organic / IG Creator / IG Ad / Snapchat / Google / Friend / Other.",
        "Approve one no-discount conversion reason and one margin-safe offer fallback.",
        "Shoot dawn + golden hour; obtain guest/staff consent; capture clean ambient sound.",
        "Edit three Reel versions from the same footage: place, product and human.",
        "Build Story templates for OPEN / SEATS / MAP / LAST CALL.",
        "Invite 5–8 suitable local creators; secure 2–3 visits, not one large expensive post.",
        "Brief service team on the advertised experience, hero product and tracking question.",
    ],
    numbering["bullet"],
)

section_break(doc)

# Media proposal
doc.add_heading("8. Media production proposal", level=1)
add_callout(
    doc,
    "Creative direction",
    "Premium naturalism: recognisable Salalah light, real mist, real condensation, real portions, real people. Editorial framing for the feed; immediate phone-native truth for Stories and Snapchat.",
)
doc.add_heading("One-week asset package", level=2)
add_table(
    doc,
    ["Asset", "Quantity", "Length / format", "Purpose"],
    [
        ("Hero vertical films", "3", "12–20 sec, 9:16", "Place / signature / people; suitable for Collab and paid."),
        ("Fast discovery cuts", "6", "5–9 sec, 9:16", "Hooks, Snap Ads, Reels testing, Spotlight."),
        ("Story/Snap clips", "35–50", "3–7 sec raw clips", "Daily live narrative and utility."),
        ("Premium stills", "18–24", "4:5 + 9:16 crops", "Carousels, covers, Maps, press/creator kit."),
        ("UGC-style testimonials", "6", "6–15 sec", "Trust, objections, creator distribution."),
        ("Route/utility clips", "6", "5–10 sec", "Landmark, entrance, seating, hours, menu, parking."),
        ("Audio", "10 clips", "5–20 sec clean WAV/M4A", "Waves, pour, ice, steam, chatter; sensory edits."),
    ],
    [2200, 1050, 2250, 3860],
)

doc.add_heading("Priority shot list", level=2)
add_table(
    doc,
    ["Priority", "Shot", "Coverage", "Proof required"],
    [
        ("A", "Beachfront reveal", "Walk-in wide; static wide; guest POV", "Haven/place recognisable in <2 sec."),
        ("A", "Seat facing sea", "Empty invitation + occupied human", "Comfort, cleanliness, spacing, true view."),
        ("A", "Passion-fruit cold brew", "Full portrait; build; condensation macro", "Actual cup, garnish, colour and 2.4 OMR price."),
        ("A", "Spanish latte", "Milk/coffee transition; handoff", "True portion and branding."),
        ("A", "Brownie", "Full serving; break/crumb macro", "Moisture, scale, actual plate."),
        ("A", "Guest arrival/payoff", "Approach → sit → first sip", "Natural expression, consent."),
        ("B", "Founder/barista", "Direct-to-camera + making ritual", "Why the signature exists; no recipe secrets."),
        ("B", "Khareef atmosphere", "Mist, palms, waves, fabric movement", "Date/time logged; no fake weather."),
        ("B", "Route/entrance", "Landmark, turn, exterior/night sign", "Usable for first-time visitor."),
        ("B", "Service proof", "Greeting, clean table, fast handoff", "Warmth and competence."),
        ("C", "Late-night mood", "Warm lights, desserts, groups", "Accurate hours and safe ambience."),
        ("C", "Texture library", "Ice, foam, steam, paper, logo, sand", "Short transition material."),
    ],
    [900, 2300, 2800, 3360],
)

doc.add_heading("Style rules", level=2)
add_bullets(
    doc,
    [
        "Light: bright directional daylight or real golden/blue hour. Avoid anonymous dark studio lighting.",
        "Colour: believable coffee, cream, matcha and karkade; cool Salalah environment; no heavy orange cast.",
        "Composition: one protagonist per shot. Alternate wide place proof, human-scale ritual and macro appetite.",
        "Motion: slow push, hand-held POV, real action. Avoid excessive speed ramps, fake zooms and template transitions.",
        "People: hands and faces create warmth; show mixed group occasions respectfully; never film identifiable guests without consent.",
        "Sound: retain waves, ice, steam and voices. Mix music beneath the physical sound rather than replacing it.",
        "Branding: natural cup/sign/menu presence by second 3; no five-second logo intros.",
        "Text: large bilingual headline, 5–8 words per frame, high contrast, safe zone, one CTA.",
        "Truth: do not generate weather, crowds, product portions or views. AI may clean distractions only if the result remains faithful.",
    ],
    numbering["bullet"],
)

doc.add_heading("Production schedule", level=2)
add_steps(
    doc,
    [
        ("Pre-production—2 hours", "Confirm operational facts; prepare clean serviceware, hero ingredients, release forms and shot order."),
        ("Dawn/morning—90 minutes", "Empty-place reveal, calm seating, route, product portraits and clean audio."),
        ("Service rehearsal—60 minutes", "Barista rituals, product builds, staff greeting and macro details."),
        ("Golden hour/peak—120 minutes", "Real guests/creators, arrivals, group plan, changing light and social proof."),
        ("Night—45 minutes", "Signage, warm tables, dessert/coffee, last-order utility."),
        ("Edit—same day", "Three hero cuts, six fast cuts, Story selects, covers and captions. Publish one while the moment is current."),
    ],
    numbering["number"],
)

section_break(doc)

# Creative scripts
doc.add_heading("9. Ready-to-shoot creative concepts", level=1)
doc.add_heading("Concept A — “Your Khareef seat” (12 sec)", level=2)
add_para(doc, "0–1s: misty sea and chair. Text: “مقعدك في الخريف هنا / Your Khareef seat is here.”")
add_para(doc, "1–5s: guest walks in; cold brew lands; condensation macro.")
add_para(doc, "5–9s: first sip, friend joins, wide view.")
add_para(doc, "9–12s: “Haven • Osara, Salalah • open today • map in bio.”")

doc.add_heading("Concept B — “Not just the view” (9 sec)", level=2)
add_para(doc, "0–2s: perfect sea view. Text: “You would come for this…”")
add_para(doc, "2–7s: cold brew build + brownie break. Text: “…you return for this.”")
add_para(doc, "7–9s: full serving and price; “Send this to your Khareef person.”")

doc.add_heading("Concept C — “Founding barista” (20–30 sec)", level=2)
add_para(doc, "Direct-to-camera: why passion fruit was chosen for the cold brew, intercut with the actual build and guest reaction. Speak naturally in Arabic; subtitle both languages. End: “Ask for the Haven Signature.”")

doc.add_heading("Concept D — “Tonight, honestly” (6 sec Snap/Story)", level=2)
add_para(doc, "One live wide shot, seat-status label, accurate hours and location. No cinematic edit. This is operational truth that converts.")

doc.add_heading("Concept E — “Arrive at Haven” (two edits from one shoot)", level=2)
add_callout(
    doc,
    "Why this matters now",
    "Until the Osara roadside sign light is repaired, the route video becomes digital wayfinding. It should remove uncertainty at the exit, car park and gate while presenting the Waterfront and Chill-out Garden as two intentional Haven experiences.",
    fill=SAND_2,
)
add_para(doc, "PRIMARY CLIPPED CUT — 18–24 seconds", size=10.3, color=NAVY)
add_steps(
    doc,
    [
        ("0–2 sec: hook", "Road approach. Text: “Missed the Haven turn before? Save this route.”"),
        ("2–6 sec: exit", "Show the recognisable approach and exact exit with a clean animated arrow. Do not depend on the unlit sign."),
        ("6–9 sec: car park", "Show the entrance and the correct place to park; keep the landmark visible."),
        ("9–13 sec: front gate", "Cut to the short walk through the gate at normal speed so first-time guests recognise it."),
        ("13–18 sec: choose the mood", "A simple split: “Waterfront — sea view & energy” / “Chill-out Garden — calm & conversation.”"),
        ("18–22 sec: payoff", "Show both areas working well, with the Waterfront first because it currently attracts more visitors."),
        ("22–24 sec: action", "“Haven • Osara, Salalah • Save this route • Maps in bio.”"),
    ],
    numbering["number"],
)
add_para(doc, "SPEED-RAMP / HYPERLAPSE CUT — 25–30 seconds", size=10.3, color=NAVY)
add_para(
    doc,
    "Use one continuous-feeling journey from the road approach to the water: 4×–8× speed on straight travel, normal speed at the exit, car-park turn, gate and final reveal. Produce a second ending that turns into the Chill-out Garden. This edit builds atmosphere; the clipped cut remains the clearer conversion asset.",
)
add_bullets(
    doc,
    [
        ("Capture once, publish four ways: ", "24-sec Instagram Reel; 12-sec ad/Snap cut; 3-frame Story route; 25–30-sec hyperlapse."),
        ("Shoot safely: ", "Use a passenger, fixed dash mount or parked camera. The driver must never hold or operate the phone."),
        ("Shoot in daylight/twilight: ", "Make the exit legible despite the broken roadside light. Do not brighten the sign in a way that implies it works."),
        ("Keep decision points slow: ", "Viewers must be able to recognise the exit, parking entrance, gate and two-area choice."),
        ("Privacy: ", "Avoid readable number plates and identifiable guests without consent."),
        ("Pin it: ", "Make this one of the three pinned Instagram posts and save the route frames to LOCATION on both Instagram and Snapchat."),
    ],
    numbering["bullet"],
)
add_para(
    doc,
    "Commercial baseline: the night of 24 July 2026 closed at approximately 110 OMR with both areas functioning well. For the next comparable night, record area chosen, “how did you find us?”, map taps and revenue. A practical first threshold is 125 OMR (+13.6%) while maintaining service quality; treat this as a test target, not proof of causation.",
)

doc.add_heading("Bilingual caption example", level=2)
add_callout(
    doc,
    "Caption",
    "الخريف أحلى على البحر. قهوتك، مقعدك، والضباب قدامك. هيفن — أوسارا، صلالة. أرسلها للشخص اللي بيجي معك اليوم.\n\nKhareef is better by the sea. Your coffee, your seat, the mist in front of you. Haven — Osara, Salalah. Send this to who is coming with you today.",
    fill=FOG,
)

doc.add_heading("Premium conversion offers: choose one only after margin approval", level=2)
add_bullets(
    doc,
    [
        ("Recommended—signature pairing: ", "Haven Passion Fruit Cold Brew + Brownie at a clear approved bundle price. It reinforces two hero products without cheapening the brand."),
        ("Experience add-on: ", "A small complimentary tasting sip with a full signature order during one quiet hour; capacity controlled."),
        ("No-discount reason: ", "A timed sunset pour, founder story or limited seating moment. Scarcity must be real, not manufactured."),
        ("Avoid: ", "Blanket percentage discounts, complicated giveaways, “tag 10 friends,” or an offer the service team cannot explain in one sentence."),
    ],
    numbering["bullet"],
)

section_break(doc)

# Measurement
doc.add_heading("10. Measurement and decision rules", level=1)
doc.add_heading("Install the measurement chain", level=2)
add_steps(
    doc,
    [
        ("Platform", "Record reach, non-follower reach, 3-second hold, average watch time, completion, shares/sends, saves, profile visits and follows."),
        ("Intent", "Record map clicks, WhatsApp taps, calls, menu views, DMs asking directions/hours/seating and Story sticker taps."),
        ("Arrival", "At POS, ask one neutral question: “How did you hear about Haven?” Use one-tap source options."),
        ("Offer", "Use unique creator/ad codes or a campaign POS item; count redemptions and average order value."),
        ("Experience", "Track wait time, stock-outs, seating complaints, review volume/rating and guest content."),
    ],
    numbering["number"],
)

add_table(
    doc,
    ["Metric", "Formula", "What it diagnoses"],
    [
        ("Share rate", "Shares ÷ reach", "Whether the post became a plan people sent."),
        ("Qualified profile rate", "Profile visits ÷ local reach", "Whether desire translated into research."),
        ("Intent rate", "(Map + WhatsApp + calls) ÷ profile visits", "Whether the profile removed friction."),
        ("Measured visit rate", "Tracked visits ÷ intent actions", "Whether intent became arrival."),
        ("Creator cost/visit", "Creator cost ÷ creator-attributed visits", "Commercial value, not vanity reach."),
        ("Campaign ROAS proxy", "Attributed gross profit ÷ media+creator cost", "Whether paid distribution earned its cost."),
    ],
    [1850, 2600, 4910],
)

doc.add_heading("Daily scorecard", level=2)
add_table(
    doc,
    ["Date/post", "Reach", "Hold / watch", "Shares", "Profile", "Map/DM", "Visits", "Revenue"],
    [
        ("", "", "", "", "", "", "", ""),
        ("", "", "", "", "", "", "", ""),
        ("", "", "", "", "", "", "", ""),
        ("", "", "", "", "", "", "", ""),
        ("", "", "", "", "", "", "", ""),
    ],
    [1500, 850, 1450, 900, 900, 1100, 1100, 1560],
    header_fill=NAVY_2,
)

doc.add_heading("Decision rules after 24–48 hours", level=2)
add_bullets(
    doc,
    [
        "High watch + high shares + low profile actions: strengthen the final frame, geotag, bio and map instruction.",
        "High profile actions + low map/DM: profile has friction or missing operational facts.",
        "High intent + low measured visits: check distance, availability, service promise, offer clarity and tracking accuracy.",
        "High views + low shares: visually pleasant but not socially useful; rewrite around a group plan or specific local insight.",
        "Low first-second hold: replace the opening; do not waste budget trying to rescue a weak hook.",
        "Strong organic qualified signals: promote that exact post, then test only one variable at a time.",
        "Crowding/service pressure rises: reduce reach spend, publish honest seat status and protect the guest experience.",
    ],
    numbering["bullet"],
)

# Values
doc.add_heading("11. What people value—and what it means for Haven", level=1)
doc.add_heading("Worldwide evidence", level=2)
add_para(
    doc,
    "Restaurant-choice research repeatedly groups the core value drivers as the “big four”: food, service, atmosphere, and price/value. Food quality is most often the leading attribute, though service and atmosphere can dominate in particular settings. Current social research and platform studies add authenticity, usefulness, entertainment and human connection. In other words, aesthetic attention earns the visit only when product and service justify it.",
)
add_table(
    doc,
    ["What people value", "Global meaning", "Haven action"],
    [
        ("Food/product quality", "The experience cannot compensate for a disappointing order.", "Lead with one signature but prove several excellent products; show full portion and texture."),
        ("Atmosphere and escape", "People pay for how a place makes time feel.", "Own the Khareef beachfront moment and daypart rituals."),
        ("Service and ease", "Warmth, speed and reliability turn first visits into advocacy.", "Market only what operations can deliver; train the team on the campaign promise."),
        ("Value", "Value is the total exchange, not only low price.", "Show price with product/place proof; use bundles, not indiscriminate discounts."),
        ("Authenticity", "Real people and real context build trust in an over-polished feed.", "Use UGC, staff voices, ambient sound and truthful weather."),
        ("Belonging/shareability", "People use hospitality to connect and signal identity.", "Write content as a group plan; create recognisable guest photo moments."),
        ("Utility", "People want fast answers before they commit.", "Hours, map, parking, seating and menu must be one tap away."),
    ],
    [1900, 3300, 4160],
)

doc.add_heading("Local evidence from Salalah/Dhofar", level=2)
add_para(
    doc,
    "Public coverage and customer-review themes around comparable Salalah cafés repeatedly praise beachfront views, relaxation, desserts/coffee, early-morning calm, sunset timing, family suitability and kind service. The strongest complaints concern waiting, crowded seating, slow or inattentive service, and a mismatch between the beautiful setting and the product/service delivered. Khareef demand is heavily Omani and GCC, with family and group travel central to the occasion.",
)
add_callout(
    doc,
    "The local conclusion",
    "People do not want a photograph of luxury. They want a beautiful, comfortable, easy outing that feels worth the drive and safe to recommend to family or friends.",
    fill=SAND,
)

doc.add_heading("Therefore Haven should", level=2)
add_bullets(
    doc,
    [
        "Sell the complete occasion: sea + seat + company + product + service.",
        "Use Arabic-first planning language and bilingual proof.",
        "Create sendable group-plan content rather than isolated glamour shots.",
        "Show prices selectively so premium feels transparent, not uncertain.",
        "Publish live seating and hours truth; never manufacture scarcity or crowds.",
        "Make the signature cold brew a memory device, then prove the rest of the menu.",
        "Treat every guest as a potential distributor: serve well, invite a photo/review, and respond quickly.",
        "Optimise for measured local visits and repeat intent; use reach as an input, not the goal.",
    ],
    numbering["bullet"],
)

# Sources
doc.add_heading("12. Sources and evidence notes", level=1)
add_para(
    doc,
    "Accessed 25 July 2026 unless otherwise stated. Platform audience numbers are ad-planning estimates and should not be interpreted as unique active users. Competitor observations are public-profile snapshots; no private analytics were available.",
)
add_source(doc, "DataReportal — Digital 2025: Oman", "https://datareportal.com/reports/digital-2025-oman", "Internet, Instagram, Snapchat, TikTok and YouTube ad-reach figures.")
add_source(doc, "Oman News Agency — Khareef Dhofar 2025 visitors", "https://omannews.gov.om/topics/en/80/show/124030", "827,115 visitors through 15 August and visitor-origin figures.")
add_source(doc, "Oman Ministry of Heritage and Tourism — Khareef 2024", "https://mht.gov.om/media-center/news/khareef-season-in-dhofar-attracted-1-048-000-visitors-in-2024/", "Full-season visitor volume and tourism-promotion context.")
add_source(doc, "Instagram / Meta — Best Practices education hub", "https://about.fb.com/news/2024/10/best-practices-education-hub-creators-instagram/", "Official guidance areas for creation, engagement, reach and guidelines.")
add_source(doc, "Instagram Help — Recommendation eligibility", "https://www.facebook.com/help/instagram/653964212890722", "Recommendation surfaces and account-status checks.")
add_source(doc, "Meta for Business — Reels Ads", "https://www.facebook.com/business/ads/facebook-instagram-reels-ads", "9:16, audio, safe zones, A/B testing and cited performance benchmarks.")
add_source(doc, "Meta for Business — Awareness objective", "https://www.facebook.com/business/ads/ad-objectives/awareness", "First seconds, sound-off clarity, local targeting and CTA guidance.")
add_source(doc, "Snapchat for Business — Creative best practices", "https://forbusiness.snapchat.com/blog/creative-best-practices-snapchat-for-business", "5–6 second native feel, product/location, urgency and sound-on figures.")
add_source(doc, "Snapchat — Modern Brand Guide", "https://forbusiness.snapchat.com/blog/modern-brand-guide-scrolling-content-snapchat", "UGC, creators, AR, edutainment and attention findings.")
add_source(doc, "Snapchat — Ad formats and specifications", "https://forbusiness.snapchat.com/advertising/ad-formats", "Single Video, Story, Commercial, AR formats and technical requirements.")
add_source(doc, "Google Business Profile", "https://www.google.com/intl/en/business/", "Search/Maps discovery, offerings, reviews and customer connection.")
add_source(doc, "Google Business Profile Help — Review requests", "https://support.google.com/business/answer/3474122?hl=en", "Compliant review-link and QR practices.")
add_source(doc, "Condé Nast Traveller Middle East — Salalah cafés", "https://www.cntravellerme.com/story/best-restaurants-and-cafes-in-salalah", "Voliere, Lantana and local experience positioning.")
add_source(doc, "Voliere Instagram", "https://www.instagram.com/voliere.om/", "Public profile snapshot: 17.1K followers; exact location/hours; own and collaborator Reels.")
add_source(doc, "Lantana Instagram", "https://www.instagram.com/lantanaom/", "Public profile snapshot: 8.2K followers; beach-led content and Menu/Location/Timing Highlights.")
add_source(doc, "55 Coffee Instagram", "https://www.instagram.com/55_coffee/", "Public profile snapshot: 62.2K followers; branch routing, product Reels and carousels.")
add_source(doc, "SevenRooms — 2025 UAE Restaurant Trends", "https://go.sevenrooms.com/rs/519-YNM-008/images/2025-UAE-Restaurant-Trends-SevenRooms.pdf?version=", "Regional evidence that social drives restaurant discovery.")
add_source(doc, "International Journal of Hospitality Management — restaurant choice review", "https://www.sciencedirect.com/science/article/abs/pii/S0963996924004393", "Food, service, atmosphere and price/value as core attributes.")
add_source(doc, "55 Coffee — brand story", "https://55coffee.co/public/index.php?page=story", "Omani identity, community and scale from Salalah.")
add_source(doc, "Independent Food Company — SALT", "https://indpt.com/", "Community, beach origin, pop-up unpredictability and signature simplicity.")
add_source(doc, "% Arabica UAE", "https://arabica.ae/", "Minimal brand system, rituals, questions and direct ordering.")
add_source(doc, "DRIFT Beach Dubai", "https://driftbeachdubai.com/", "Destination packaging, dayparts, booking and seasonal operations.")
add_source(doc, "Koko Bay", "https://kokobay.co/", "Beachfront occasion, offers and reservation path.")

doc.add_heading("Internal Haven materials used", level=2)
add_bullets(
    doc,
    [
        "Current bilingual menu and prices, including Haven Passion Fruit Cold Brew (2.4 OMR), Spanish Latte and Brownie.",
        "Haven logo, existing brand imagery, and the Haven Print Menu: Editorial Art Direction Study—premium restraint, product truth, Salalah light and four-image grammar.",
    ],
    numbering["bullet"],
)

# Prevent widows/orphans and add metadata
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True

props = doc.core_properties
props.title = "The Haven Effect — Khareef Social Media Instruction Book"
props.subject = "Instagram, Snapchat and footfall strategy for Haven, Salalah"
props.author = "Haven"
props.keywords = "Haven, Salalah, Khareef, Instagram, Snapchat, café, beachfront, social media"

doc.save(OUT)
print(OUT)
